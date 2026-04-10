"""
update_docx.py — Patches docs/JsonT.docx with Privacy & Encryption content
and an Applications section.

Usage:
    pip install python-docx
    python docs/update_docx.py

The script appends new sections directly after the last existing heading it
finds in the document, so it is safe to re-run (it checks for duplicates).
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOCX_PATH = os.path.join(os.path.dirname(__file__), "JsonT.docx")
doc = Document(DOCX_PATH)

# ── helpers ───────────────────────────────────────────────────────────────────

def heading_texts(doc):
    return {p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")}

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_body(doc, text):
    doc.add_paragraph(text)

def add_code(doc, code):
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

existing = heading_texts(doc)

# ── Section: Privacy & Encryption ────────────────────────────────────────────

if "Privacy & Encryption" not in existing:
    add_heading(doc, "Privacy & Encryption", level=1)

    add_heading(doc, "Sensitive fields (~)", level=2)
    add_body(doc,
        "Prefix a scalar field type with ~ to mark it as privacy-sensitive. "
        "The field is encrypted to a base64:<b64> wire token on write and "
        "travels as an opaque Encrypted value through the pipeline."
    )
    add_code(doc,
        "Person: {\n"
        "  fields: {\n"
        "    str:  name,\n"
        "    ~str: ssn,          // encrypted on wire\n"
        "    ~str: cardNumber?   // optional, encrypted when present\n"
        "  }\n"
        "}"
    )

    add_heading(doc, "Wire format", level=2)
    add_body(doc,
        "Encrypted fields are serialised as a JSON string prefixed with base64: "
        "followed by the Base64-encoded ciphertext bytes produced by "
        "CryptoConfig.encrypt()."
    )
    add_code(doc, '"base64:SGVsbG8gV29ybGQ="')
    add_body(doc,
        "When reading, a base64:... string at a sensitive-field position is "
        "parsed directly into an Encrypted value — no decryption occurs at "
        "parse time."
    )

    add_heading(doc, "Decrypt operation (derived schemas)", level=2)
    add_body(doc,
        "The decrypt(field1, field2, ...) derived-schema operation decrypts "
        "named fields inline during transform. It requires a CryptoConfig at "
        "runtime and is idempotent — already-plaintext fields are left unchanged."
    )
    add_code(doc,
        "PersonDecrypted: FROM Person {\n"
        "  operations: (\n"
        "    decrypt(ssn, cardNumber)\n"
        "  )\n"
        "}"
    )
    add_body(doc,
        "Calling transform() without a CryptoConfig when a decrypt operation "
        "is present returns TransformError::DecryptFailed (Rust) or throws "
        "JsonTError.Transform.DecryptFailed (Java)."
    )

    add_heading(doc, "On-demand decryption API", level=2)
    add_body(doc,
        "Individual values and rows expose a decrypt API for cases where a "
        "full pipeline decrypt operation is unnecessary:"
    )
    add_bullet(doc, "value.decrypt_str(fieldName, crypto)  → Option<String> / Optional<String>")
    add_bullet(doc, "value.decrypt_bytes(fieldName, crypto) → Option<Vec<u8>> / Optional<byte[]>")
    add_bullet(doc, "row.decrypt_field_str(index, fieldName, crypto)  (Rust) → Option<String>")
    add_bullet(doc, "row.decryptField(index, fieldName, crypto)  (Java) → Optional<String>")
    add_body(doc,
        "All methods return None / Optional.empty() for non-encrypted values, "
        "making them safe to call without first inspecting the value type."
    )

    add_heading(doc, "CryptoConfig interface", level=2)
    add_body(doc,
        "CryptoConfig is a pluggable, implementation-agnostic interface. "
        "The field name is passed to both encrypt and decrypt to support "
        "per-field or per-field-type key strategies."
    )
    add_code(doc,
        "// Rust\n"
        "trait CryptoConfig {\n"
        "    fn encrypt(&self, field: &str, plaintext: &[u8]) -> Result<Vec<u8>, CryptoError>;\n"
        "    fn decrypt(&self, field: &str, ciphertext: &[u8]) -> Result<Vec<u8>, CryptoError>;\n"
        "}"
    )
    add_code(doc,
        "// Java\n"
        "interface CryptoConfig {\n"
        "    byte[] encrypt(String field, byte[] plaintext) throws CryptoError;\n"
        "    byte[] decrypt(String field, byte[] ciphertext) throws CryptoError;\n"
        "}"
    )
    add_body(doc,
        "PassthroughCryptoConfig is the built-in identity implementation "
        "(bytes in = bytes out) — useful for testing."
    )

# ── Section: Applications ─────────────────────────────────────────────────────

if "Applications" not in existing:
    add_heading(doc, "Applications", level=1)

    add_body(doc,
        "JsonT is a general-purpose data contract format. The following domains "
        "have been modelled using JsonT schemas included in this repository."
    )

    add_heading(doc, "Fintech & Payments", level=2)
    add_bullet(doc,
        "ISO 20022 (pain.001) — Payment initiation with IBAN regex validation, "
        "currency code constraints, and hierarchical party/agent structures."
    )
    add_bullet(doc,
        "NACHA / ACH — US Automated Clearing House files with record-type "
        "constant checks and fixed-length field constraints."
    )
    add_bullet(doc,
        "FIX Protocol — Trading messages with conditional requirements "
        "(e.g. ordType == \"2\" -> required(price)) and standard tag mapping."
    )

    add_heading(doc, "Healthcare", level=2)
    add_bullet(doc,
        "HL7 834 (Benefit Enrollment) — ISA, ST, BGN, INS, REF, DTP, NM1, "
        "DMG, HD segments with loop hierarchies (1000A/B, 2000, 2300) and "
        "X12 code-set enums."
    )
    add_body(doc,
        "Sensitive PII fields (SSNs, member IDs) can be marked with ~ and "
        "encrypted at the field level without changing the positional row format."
    )

    add_heading(doc, "Analytics & Data Pipelines", level=2)
    add_bullet(doc, "Batch API responses — 45–55% smaller than equivalent JSON for large record sets.")
    add_bullet(doc, "Event streams and logs — O(1) streaming validation with no memory growth.")
    add_bullet(doc, "Analytics exports — Schema-enforced types prevent downstream type coercion errors.")
    add_bullet(doc, "Data lakehouse ingestion — Derived schemas project and filter rows in-flight.")

    add_heading(doc, "Multi-system Integration", level=2)
    add_body(doc,
        "Derived schemas let a single canonical schema (e.g. Order) fan out "
        "to multiple consumers through typed projections — each consumer "
        "receives only the fields it needs, validated against the original schema."
    )
    add_code(doc,
        "Order (Straight): id, product, quantity, price, customerId, internalCost\n"
        "  └─ OrderPublic (Derived):   project(id, product, price)     → 3 fields\n"
        "  └─ OrderBilling (Derived):  exclude(internalCost)           → 5 fields\n"
        "  └─ OrderAnalytics (Derived): rename(customerId as cid), ...\n"
        "                               filter price > 0"
    )

    add_heading(doc, "Privacy-preserving Pipelines", level=2)
    add_body(doc,
        "The ~ privacy marker combined with a pluggable CryptoConfig makes it "
        "straightforward to build pipelines where sensitive fields (SSNs, card "
        "numbers, PII) travel encrypted end-to-end and are decrypted only at "
        "authorised consumers:"
    )
    add_bullet(doc, "Producer encrypts on write using RowWriter.writeRow(..., crypto, writer).")
    add_bullet(doc, "Pipeline stages validate and transform without ever seeing plaintext.")
    add_bullet(doc, "Authorised consumer decrypts via transform_with_crypto or on-demand decryptField.")

# ── Save ──────────────────────────────────────────────────────────────────────

doc.save(DOCX_PATH)
print(f"Saved: {DOCX_PATH}")
print(f"Sections added: Privacy & Encryption={'Privacy & Encryption' not in existing}, "
      f"Applications={'Applications' not in existing}")
